from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


def main() -> None:
    source = Path(sys.argv[1])
    output = Path(sys.argv[2]) if len(sys.argv) > 2 else None
    document = Document(source)
    paragraphs = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text:
            paragraphs.append(
                {
                    "index": index,
                    "style": paragraph.style.name if paragraph.style else "",
                    "text": text,
                }
            )

    tables = []
    for table_index, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"index": table_index, "rows": rows})

    result = {
        "source": str(source),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "paragraphs": paragraphs,
        "tables": tables,
    }
    serialized = json.dumps(result, ensure_ascii=False, indent=2)
    if output:
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text(serialized, encoding="utf-8")
        print(output)
    else:
        print(serialized)


if __name__ == "__main__":
    main()
